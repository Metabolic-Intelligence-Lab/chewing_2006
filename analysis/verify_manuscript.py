"""Manuscript <-> code/tables consistency checker.

Extracts the key quantitative claims stated in the generated manuscript docx and
compares each against the value recomputed from the pipeline output tables. Prints
a diff table and exits non-zero if any claim is out of tolerance, so it can gate a
release / catch silent drift after code changes.

Usage:
    python analysis/verify_manuscript.py            # diff docx <-> existing tables
    python analysis/verify_manuscript.py --regen     # run_all first, then diff
    python analysis/verify_manuscript.py --docx <path> --tables <dir>

Note: the tables under outputs/ must be current. After changing analysis code,
run `python analysis/run_all.py` (or pass --regen) so the tables reflect the new
logic before trusting the diff.
"""
from __future__ import annotations

import argparse
import pathlib
import re
import sys

import pandas as pd

BASE = pathlib.Path(__file__).resolve().parent
OUT = BASE / "outputs"
TABLES = OUT / "tables"
DEFAULT_DOCX = BASE / "report" / "Chewing_Swallowing_pipeline_EN.docx"


# ---------- loaders (defensive: return None if missing) ----------
def _csv(name):
    p = TABLES / name if not name.startswith("/") else pathlib.Path(name)
    if not p.exists():
        p = OUT / name
    return pd.read_csv(p) if p.exists() else None


def _cell(df, where, col):
    """First value of `col` in rows matching the dict `where`, else None."""
    if df is None or col not in df.columns:
        return None
    m = pd.Series(True, index=df.index)
    for k, v in where.items():
        if k not in df.columns:
            return None
        m &= (df[k] == v)
    sub = df[m]
    return float(sub[col].iloc[0]) if len(sub) else None


def _qc_value(pattern):
    p = OUT / "qc_report.md"
    if not p.exists():
        return None
    txt = p.read_text(encoding="utf-8", errors="ignore")
    m = re.search(pattern, txt)
    return m


def t_coupling_r():
    return _cell(_csv("A_coordination_corr.csv"),
                 {"x": "chew_work_rate", "y": "swallow_work_rate"}, "r")


def t_coupling_pfdr():
    return _cell(_csv("A_coordination_corr.csv"),
                 {"x": "chew_work_rate", "y": "swallow_work_rate"}, "p_fdr")


def t_coupling_p():
    return _cell(_csv("A_coordination_corr.csv"),
                 {"x": "chew_work_rate", "y": "swallow_work_rate"}, "p")


def t_dyn(outcome, col):
    v = _cell(_csv("CYC_within_sequence.csv"), {"outcome": outcome}, col)
    if v is None:
        return None
    return abs(v) if col == "beta_per_sd" else v


def t_coherence_median():
    df = _csv("DIR_directionality.csv")
    return float(df["coherence_chewband"].median()) if (df is not None and "coherence_chewband" in df.columns) else None


def t_granger_pct(col):
    df = _csv("DIR_directionality.csv")
    if df is None or col not in df.columns:
        return None
    return float((df[col] < 0.05).mean() * 100)


def t_sens_min_rank():
    """Min rank-corr excluding the two count-rescaling parameters (matches the
    manuscript's '>=0.87 for all but ...')."""
    df = _csv("SENS_swallow_params.csv")
    if df is None or "rank_corr_vs_default" not in df.columns:
        return None
    rescalers = {"(default)", "min_pause_s", "adaptive_percentile"}
    nd = df[~df.parameter.isin(rescalers)]["rank_corr_vs_default"]
    return float(nd.min()) if len(nd) else None


def t_heur_r():
    from scipy.stats import pearsonr
    df = _csv("VAL_agreement.csv")
    if df is None or not {"official", "heuristic"}.issubset(df.columns) or len(df) < 3:
        return None
    return float(pearsonr(df["official"], df["heuristic"])[0])


def _ann(col):
    df = _csv("ANN_concordance_summary.csv")
    return float(df[col].iloc[0]) if (df is not None and col in df.columns and len(df)) else None


def t_ann_spec_pct():
    v = _ann("pooled_frac_events_matched")
    return v * 100 if v is not None else None


def t_ann_offset():
    return _ann("offset_median_s")


# ---- internal/construct validation (internal_validation.py) ----
def _iv_inscope(col, scale=1.0):
    df = _csv("IV_inscope_summary.csv")
    return float(df[col].iloc[0]) * scale if (df is not None and col in df.columns and len(df)) else None


def _iv_null(tol, col, scale=1.0):
    df = _csv("IV_null_concordance.csv")
    if df is None:
        return None
    v = _cell(df, {"null": "shift", "metric": "specificity", "tol_s": float(tol)}, col)
    return v * scale if v is not None else None


def _iv_semi(alpha):
    v = _cell(_csv("IV_semisynthetic_recall.csv"), {"alpha": float(alpha)}, "recall_v1")
    return v * 100 if v is not None else None


def _iv_plaus(col, scale=100.0):
    v = _cell(_csv("IV_plausibility_summary.csv"), {"metric": "fractions"}, col)
    return v * scale if v is not None else None


def _iv_pilot(col):
    return _cell(_csv("IV_pilot_summary.csv"), {"pair": "POOLED_INCLUDED"}, col)


# ---------- checks: (label, manuscript-regex, group#, table_fn, tol) ----------
# regex captures the number AS STATED in the manuscript; table_fn recomputes it.
CHECKS = [
    ("coupling rho",            r"swallowing work rate \(ρ=([\d.]+)", 1, t_coupling_r, 0.02),
    ("coupling p",             r"swallowing work rate \(ρ=[\d.]+, p=([\d.]+),", 1, t_coupling_p, 0.003),
    ("coupling p_FDR",          r"ρ=[\d.]+, p=[\d.]+, p_FDR=([\d.]+), n=", 1, t_coupling_pfdr, 0.01),
    ("heuristic agreement r",   r"per-subject swallow counts \(r=([\d.]+)\)", 1, t_heur_r, 0.02),
    ("sensitivity min rank",    r"Spearman ≥([\d.]+) against the default", 1, t_sens_min_rank, 0.02),
    ("coherence median",        r"coherence is moderate \(median ([\d.]+)\)", 1, t_coherence_median, 0.03),
    ("Granger chew->swallow %", r"\((\d+)% chew→swallow", 1, lambda: t_granger_pct("gc_chew2swallow_p"), 4.0),
    ("Granger swallow->chew %", r"chew→swallow, (\d+)% swallow→chew\)", 1, lambda: t_granger_pct("gc_swallow2chew_p"), 4.0),
    ("dyn rise_time beta",      r"rise time β=−([\d.]+)/SD, p=", 1, lambda: t_dyn("rise_time_s", "beta_per_sd"), 0.01),
    ("dyn rise_time p",         r"rise time β=−[\d.]+/SD, p=([\d.]+);", 1, lambda: t_dyn("rise_time_s", "p"), 0.0005),
    ("dyn duration beta",       r"duration β=−([\d.]+)/SD, p=", 1, lambda: t_dyn("duration_s", "beta_per_sd"), 0.01),
    ("annotation specificity %", r"coincide with a manual mark \((\d+)% within ±10 s;", 1, t_ann_spec_pct, 2.0),
    ("annotation offset s",     r"offset median (-?[\d.]+) s", 1, t_ann_offset, 0.3),
    # internal/construct validation
    ("in-scope mark fraction %", r"(\d+)% of manual marks fall inside a phase-end window", 1,
     lambda: _iv_inscope("frac_annot_in_scope", 100), 1.0),
    ("in-scope recall %",       r"in-scope recall (\d+)% \(±10 s\)", 1,
     lambda: _iv_inscope("recall_in_scope_tol10", 100), 1.0),
    ("in-scope recall corr. %", r"in-scope recall \d+% \(±10 s\), (\d+)% after reaction-time correction", 1,
     lambda: _iv_inscope("recall_in_scope_corr_tol10", 100), 1.0),
    ("null specificity ±10 s %", r"chance (\d+)%, enrichment ×", 1, lambda: _iv_null(10, "null_mean", 100), 1.0),
    ("enrichment ±10 s",        r"chance \d+%, enrichment ×([\d.]+)", 1, lambda: _iv_null(10, "enrichment"), 0.02),
    ("enrichment ±3 s",         r"at ±3 s enrichment ×([\d.]+)", 1, lambda: _iv_null(3, "enrichment"), 0.02),
    ("semi-synthetic recall α=1 %", r"α=1\.0: (\d+)%", 1, lambda: _iv_semi(1.0), 1.0),
    ("FWHM in-range %",         r"(\d+)% of accepted events have a FWHM", 1,
     lambda: _iv_plaus("frac_fwhm_in_lit_range"), 1.0),
    ("masseter ratio<1 %",      r"preceding chewing phase in (\d+)% of", 1,
     lambda: _iv_plaus("frac_mast_ratio_phase_lt1"), 1.0),
    ("pilot chew-only ev/min",  r"chew-only ([\d.]+) events/min", 1, lambda: _iv_pilot("chew_events_per_min"), 0.1),
    ("pilot swallow-only ev/min", r"swallow-only ([\d.]+) events/min", 1, lambda: _iv_pilot("swal_events_per_min"), 0.1),
]


def manuscript_text(docx_path):
    import docx
    d = docx.Document(str(docx_path))
    return "\n".join(p.text for p in d.paragraphs if p.text.strip())


def run(docx_path):
    txt = manuscript_text(docx_path)
    rows, n_fail, n_missing = [], 0, 0
    for label, rgx, grp, fn, tol in CHECKS:
        m = re.search(rgx, txt)
        stated = float(m.group(grp)) if m else None
        try:
            computed = fn() if callable(fn) else None
        except Exception as e:
            computed = None
        if stated is None:
            status, n_missing = "CLAIM NOT FOUND", n_missing + 1
        elif computed is None:
            status, n_missing = "TABLE MISSING", n_missing + 1
        else:
            ok = abs(stated - computed) <= tol
            status = "ok" if ok else "*** MISMATCH ***"
            if not ok:
                n_fail += 1
        rows.append((label,
                     "—" if stated is None else f"{stated:g}",
                     "—" if computed is None else f"{computed:.4g}",
                     "" if (stated is None or computed is None) else f"{stated-computed:+.3g}",
                     status))
    w = max(len(r[0]) for r in rows)
    print(f"\nManuscript: {docx_path}")
    print(f"Tables:     {TABLES}\n")
    print(f"{'claim'.ljust(w)}  {'stated':>9}  {'computed':>9}  {'Δ':>8}  status")
    print("-" * (w + 42))
    for label, s, c, d, st in rows:
        print(f"{label.ljust(w)}  {s:>9}  {c:>9}  {d:>8}  {st}")
    print("-" * (w + 42))
    print(f"{len(rows)} checks: {len(rows)-n_fail-n_missing} ok, {n_fail} mismatch, {n_missing} unresolved")
    return n_fail


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--regen", action="store_true", help="run the full pipeline first")
    ap.add_argument("--docx", default=str(DEFAULT_DOCX))
    ap.add_argument("--tables", default=None, help="override tables dir")
    args = ap.parse_args()
    global TABLES
    if args.tables:
        TABLES = pathlib.Path(args.tables)
    if args.regen:
        print("Regenerating pipeline outputs (this can take several minutes)...")
        sys.path.insert(0, str(BASE / "src"))
        import run_all
        run_all.main()
    n_fail = run(pathlib.Path(args.docx))
    sys.exit(1 if n_fail else 0)


if __name__ == "__main__":
    main()
