"""Shared python-docx helpers (headings, paragraphs, figures, tables, number formatting)
and the cohort demographics table used by make_article.py.
"""
from __future__ import annotations

import pathlib
import re

import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

PAGE_TEXT_WIDTH_IN = 6.5   # Letter, 1-inch margins -> usable text width


def _nd(text):
    """Strip em dashes (U+2014) from rendered text, substituting commas (or
    tidying punctuation), so the manuscript reads without em dashes. En dashes
    (U+2013, used for numeric ranges) and the minus sign (U+2212) are left as is.
    """
    if not isinstance(text, str) or "—" not in text:
        return text
    s = text.replace(" — ", ", ")
    s = s.replace("— ", ", ").replace(" —", ", ").replace("—", ", ")
    # tidy any doubled/misplaced punctuation introduced by the substitution
    s = re.sub(r",\s*,", ", ", s)
    s = re.sub(r"([;:.])\s*,\s*", r"\1 ", s)
    s = re.sub(r"\s+,", ",", s)
    s = re.sub(r",\s*\)", ")", s)
    s = re.sub(r"\(\s*,\s*", "(", s)
    return s

OUT = pathlib.Path(__file__).resolve().parents[1] / "outputs"
FIG = OUT / "figures"
TAB = OUT / "tables"
REPORT = pathlib.Path(__file__).resolve().parents[1] / "report"
CENTER = WD_ALIGN_PARAGRAPH.CENTER


def h(doc, text, level=1):
    p = doc.add_heading(_nd(text), level=level)
    for r in p.runs:
        r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    return p


def para(doc, text, italic=False, align=None, size=11):
    p = doc.add_paragraph()
    r = p.add_run(_nd(text)); r.italic = italic; r.font.size = Pt(size)
    if align:
        p.alignment = align
    return p


def caption(doc, label, text):
    p = doc.add_paragraph()
    r = p.add_run(_nd(label) + " "); r.bold = True; r.font.size = Pt(9)
    r2 = p.add_run(_nd(text)); r2.font.size = Pt(9); r2.italic = True
    p.alignment = CENTER


def add_figure(doc, path, width_in=6.2):
    if pathlib.Path(path).exists():
        # never exceed the usable text width, else the image overflows the margin
        doc.add_picture(str(path), width=Inches(min(width_in, PAGE_TEXT_WIDTH_IN)))
        doc.paragraphs[-1].alignment = CENTER


def f3(v):
    try: return f"{float(v):.3f}"
    except Exception: return "-"


def f2(v):
    try: return f"{float(v):.2f}"
    except Exception: return "-"


def fp(v):
    try:
        v = float(v); return "<0.001" if v < 1e-3 else f"{v:.3f}"
    except Exception: return "-"


def add_table(doc, df, headers, fmt=None, bold_first=False):
    fmt = fmt or {}
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, hd in enumerate(headers):
        c = t.rows[0].cells[j]; c.text = _nd(hd)
        for p in c.paragraphs:
            for r in p.runs: r.bold = True; r.font.size = Pt(9)
    for _, row in df.iterrows():
        cells = t.add_row().cells
        for j, key in enumerate(df.columns):
            fn = fmt.get(key)
            cells[j].text = _nd(fn(row[key]) if fn else str(row[key]))
            for p in cells[j].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
                    if bold_first and j == 0:
                        r.bold = True
    return t


# ---------- demographics (Table 1) ----------
def _yes(s):
    return s.astype(str).str.normalize("NFKD").str.encode("ascii", "ignore") \
            .str.decode("ascii").str.lower().str.startswith("si")


def _contin(d, col, unit="", dec=1):
    s = pd.to_numeric(d[col], errors="coerce").dropna()
    if len(s) == 0:
        return "-", "-"
    return (f"{s.median():.{dec}f} [{s.quantile(.25):.{dec}f}–{s.quantile(.75):.{dec}f}]"
            + (f" {unit}" if unit else ""), str(len(s)))


def _binary(d, col, fn):
    s = d[col].dropna()
    if len(s) == 0:
        return "-", "-"
    pos = int(fn(s).sum())
    return f"{pos} ({pos / len(s) * 100:.0f}%)", str(len(s))


def demographics_df(d):
    rows = []
    n = len(d)
    rows.append(("N (adults ≥18)", str(n), ""))
    f = int((d.sex == "Donna").sum()); m = int((d.sex == "Uomo").sum())
    rows.append(("Sex — women / men, n", f"{f} / {m}", str(f + m)))
    for label, col, unit, dec in [
        ("Age, years", "age", "", 0),
        ("Height, cm", "height_cm", "", 0),
        ("Weight, kg", "weight_kg", "", 1),
        ("BMI, kg/m²", "bmi", "", 1),
        ("Body fat, %", "fat_pct", "", 1),
        ("Muscle mass, %", "muscle_pct", "", 1),
        ("Muscle-mass index, kg/m²", "muscle_index", "", 1),
        ("Visceral fat, %", "visceral_fat", "", 1),
        ("Resting metabolism, kcal", "resting_metab", "", 0),
        ("Handgrip (max), kg", "handgrip_max", "", 1),
        ("VO₂max, mL/kg/min", "vo2max", "", 1),
    ]:
        val, nn = _contin(d, col, unit, dec)
        rows.append((label, val, nn))
    # signal features (medians)
    for label, col, dec in [
        ("Chewing phases, n", "chew_phases", 0),
        ("Number of chews, n", "chew_n", 0),
        ("Chewing work rate, a.u.", "chew_work_rate", 3),
        ("Deglutition events, n", "swallow_events", 0),
        ("Subjects with ≥1 swallow, n (%)", None, None),
    ]:
        if col is None:
            vs = int(d.valid_swallow.sum())
            rows.append((label, f"{vs} ({vs / n * 100:.0f}%)", str(n)))
        else:
            val, nn = _contin(d, col, "", dec)
            rows.append((label, val, nn))
    # clinical prevalences
    for label, col, fn in [
        ("Chewing pain (any)", "chew_pain", lambda s: s > 0),
        ("Self-reported dysphagia", "dysphagia", _yes),
        ("Night bite / bruxism", "night_byte", _yes),
        ("Occlusion problem / missing teeth", "occlusion_problem", _yes),
        ("Gastro-oesophageal reflux", "reflux", _yes),
        ("Smoker (current/occasional)", "smoker", lambda s: s.astype(str).str.lower().str.startswith("si")),
    ]:
        if col in d.columns:
            val, nn = _binary(d, col, fn)
            rows.append((label, val, nn))
    return pd.DataFrame(rows, columns=["Characteristic", "Value", "n"])


